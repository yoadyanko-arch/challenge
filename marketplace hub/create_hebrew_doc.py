from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

def set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def set_run_rtl(run):
    rPr = run._r.get_or_add_rPr()
    rtl = OxmlElement('w:rtl')
    rPr.append(rtl)

def add_heading(doc, text, level=1, color=None, size=None):
    p = doc.add_paragraph()
    set_rtl(p)
    run = p.add_run(text)
    set_run_rtl(run)
    run.bold = True
    if size:
        run.font.size = Pt(size)
    elif level == 1:
        run.font.size = Pt(28)
    elif level == 2:
        run.font.size = Pt(18)
    else:
        run.font.size = Pt(14)
    if color:
        run.font.color.rgb = RGBColor(*color)
    else:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x3E)
    run.font.name = 'David'
    return p

def add_body(doc, text, indent=False):
    p = doc.add_paragraph()
    set_rtl(p)
    run = p.add_run(text)
    set_run_rtl(run)
    run.font.size = Pt(12)
    run.font.name = 'David'
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    if indent:
        p.paragraph_format.right_indent = Cm(1)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph()
    set_rtl(p)
    run = p.add_run(f"◆  {text}")
    set_run_rtl(run)
    run.font.size = Pt(12)
    run.font.name = 'David'
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p.paragraph_format.right_indent = Cm(1)
    return p

def add_phase(doc, num, title, desc):
    p = doc.add_paragraph()
    set_rtl(p)
    run1 = p.add_run(f"שלב {num} – {title}: ")
    set_run_rtl(run1)
    run1.bold = True
    run1.font.size = Pt(12)
    run1.font.name = 'David'
    run1.font.color.rgb = RGBColor(0x8B, 0x5C, 0xF6)
    run2 = p.add_run(desc)
    set_run_rtl(run2)
    run2.font.size = Pt(12)
    run2.font.name = 'David'
    run2.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p.paragraph_format.right_indent = Cm(1)

def add_divider(doc):
    p = doc.add_paragraph()
    run = p.add_run("─" * 55)
    run.font.color.rgb = RGBColor(0xD4, 0xAF, 0x37)
    run.font.size = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_section_title(doc, number, title):
    p = doc.add_paragraph()
    set_rtl(p)
    # Number badge
    run_num = p.add_run(f"{number}. ")
    set_run_rtl(run_num)
    run_num.bold = True
    run_num.font.size = Pt(20)
    run_num.font.color.rgb = RGBColor(0xD4, 0xAF, 0x37)
    run_num.font.name = 'David'
    # Title
    run_title = p.add_run(title)
    set_run_rtl(run_title)
    run_title.bold = True
    run_title.font.size = Pt(20)
    run_title.font.color.rgb = RGBColor(0x1A, 0x1A, 0x3E)
    run_title.font.name = 'David'
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    return p

doc = Document()

# Page setup
section = doc.sections[0]
section.page_width = Cm(29.7)
section.page_height = Cm(21)
section.left_margin = Cm(3)
section.right_margin = Cm(3)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)

# Set document RTL
body = doc.element.body
sectPr = body.find(qn('w:sectPr'))
if sectPr is None:
    sectPr = OxmlElement('w:sectPr')
    body.append(sectPr)
bidi_elem = OxmlElement('w:bidi')
sectPr.append(bidi_elem)

# ─── Title Page ───
p_title = doc.add_paragraph()
set_rtl(p_title)
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(40)
p_title.paragraph_format.space_after = Pt(10)
run_title = p_title.add_run("Marketplace Hub")
run_title.bold = True
run_title.font.size = Pt(40)
run_title.font.color.rgb = RGBColor(0x1A, 0x1A, 0x3E)
run_title.font.name = 'David'

p_sub = doc.add_paragraph()
set_rtl(p_sub)
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub.paragraph_format.space_after = Pt(6)
run_sub = p_sub.add_run("אקוסיסטם הצמיחה עבור מותגים בתחילת הדרך")
set_run_rtl(run_sub)
run_sub.bold = True
run_sub.font.size = Pt(18)
run_sub.font.color.rgb = RGBColor(0xD4, 0xAF, 0x37)
run_sub.font.name = 'David'

p_line = doc.add_paragraph()
p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_line = p_line.add_run("━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
run_line.font.color.rgb = RGBColor(0xD4, 0xAF, 0x37)
run_line.font.size = Pt(10)

doc.add_paragraph()

# ─── 1. Vision ───
add_section_title(doc, 1, "חזון")
add_body(doc, "להפוך לבית ולאקוסיסטם הצמיחה המרכזי עבור הדור הבא של המותגים בתחילת הדרך בישראל.")
add_body(doc, "Marketplace Hub שואפת לספק למותגים צעירים את הכלים, החיבורים, ההזדמנויות והתמיכה שהם צריכים כדי להפוך מעסקים קטנים למותגים צומחים.")
add_divider(doc)

# ─── 2. Problem ───
add_section_title(doc, 2, "הבעיה")
add_body(doc, "מותגים רבים בתחילת הדרך מתקשים לא בגלל המוצר שלהם, אלא בגלל חוסר בתשתית עסקית נכונה לצמיחה.")
add_divider(doc)

# ─── 3. Solution ───
add_section_title(doc, 3, "הפתרון")
add_body(doc, "Marketplace Hub היא פלטפורמה שנועדה להעניק למותגים בתחילת הדרך מעטפת צמיחה מלאה הכוללת:")
add_bullet(doc, "תשתית עסקית")
add_bullet(doc, "שירותים מקצועיים")
add_bullet(doc, "שיווק")
add_bullet(doc, "הזדמנויות ריטייל")
add_bullet(doc, "שותפויות ונטוורקינג")
add_bullet(doc, "כלי צמיחה")
add_divider(doc)

# ─── 4. Ecosystem ───
add_section_title(doc, 4, "מבנה האקוסיסטם")

layers = [
    ("Marketplace", "שכבת החשיפה והקהילה"),
    ("Marketplace Hub", "שכבת הצמיחה והתשתית"),
    ("Partners", "שותפים מקצועיים וספקי שירות"),
    ("Brands", "מותגים בתחילת הדרך המשתמשים באקוסיסטם לצמיחה"),
]
for eng, heb in layers:
    p = doc.add_paragraph()
    set_rtl(p)
    run_heb = p.add_run(f"{heb}  |  ")
    set_run_rtl(run_heb)
    run_heb.font.size = Pt(12)
    run_heb.font.name = 'David'
    run_heb.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    run_eng = p.add_run(eng)
    run_eng.bold = True
    run_eng.font.size = Pt(12)
    run_eng.font.name = 'David'
    run_eng.font.color.rgb = RGBColor(0x8B, 0x5C, 0xF6)
    p.paragraph_format.right_indent = Cm(1)

add_divider(doc)

# ─── 5. Services ───
add_section_title(doc, 5, "סוגי שירותים")
services = [
    "ליווי עסקי ומשפטי",
    "שיווק ומיתוג",
    "חשיפה והזדמנויות צמיחה",
    "גישה לריטייל",
    "קהילה ונטוורקינג",
]
for s in services:
    add_bullet(doc, s)
add_divider(doc)

# ─── 6. GTM ───
add_section_title(doc, 6, "מודל פריצה לשוק")
gtm_phases = [
    ("1", "השקה", "דרך האקוסיסטם הקיים של Marketplace"),
    ("2", "בניית שירותי צמיחה", "ושותפויות אסטרטגיות"),
    ("3", "בניית קהילה", "ותוכן מקצועי"),
    ("4", "הרחבה", "לפלטפורמה scalable ומודל Membership"),
]
for num, title, desc in gtm_phases:
    add_phase(doc, num, title, desc)
add_divider(doc)

# ─── 7. Business Model ───
add_section_title(doc, 7, "מודל עסקי")
biz_items = [
    "חבילות צמיחה",
    "עמלות משותפים ושותפויות",
    "מודל Membership",
    "אירועים וסדנאות",
    "גישה לריטייל",
    "הכנסות עתידיות מפלטפורמה וכלי SaaS",
]
for b in biz_items:
    add_bullet(doc, b)
add_divider(doc)

# ─── 8. Roadmap ───
add_section_title(doc, 8, "Roadmap לצמיחה")
roadmap_phases = [
    ("1", "בניית הבסיס", "רשת השותפים"),
    ("2", "חיזוק הקהילה", "ושיתופי הפעולה"),
    ("3", "בניית תשתית scalable", "ומודל Membership"),
    ("4", "הרחבה לכלים דיגיטליים", "ותשתית פלטפורמה"),
]
for num, title, desc in roadmap_phases:
    add_phase(doc, num, title, desc)
add_divider(doc)

# ─── 9. Long-Term Goal ───
add_section_title(doc, 9, "מטרה ארוכת טווח")
add_body(doc, "להפוך לאקוסיסטם הצמיחה המוביל בישראל עבור מותגים בתחילת הדרך ולהעניק למותגים את כל מה שהם צריכים כדי לצמוח.")
doc.add_paragraph()

# Footer-like closing
p_close = doc.add_paragraph()
p_close.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_close = p_close.add_run("Marketplace Hub  ·  2026")
run_close.font.size = Pt(10)
run_close.font.color.rgb = RGBColor(0xD4, 0xAF, 0x37)
run_close.font.name = 'David'

output_path = "/Users/yoadyankowitz/Library/CloudStorage/OneDrive-ReichmanUniversity/יועד/כללי/marketplace hub/Marketplace_Hub_Hebrew_V2.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
